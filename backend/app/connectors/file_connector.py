"""
File connector — handles local file uploads (PDF, TXT, DOCX).

Extracts text content and returns NormalizedDocument instances.
"""

import os
import uuid
import logging
from typing import List

from pypdf import PdfReader
from docx import Document as DocxDocument

from app.connectors.base import BaseConnector, NormalizedDocument

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


class FileConnector(BaseConnector):
    """Connector for local file uploads."""

    def __init__(self):
        self._file_path: str | None = None
        self._filename: str | None = None

    async def connect(self, credentials: dict) -> bool:
        """
        For files, 'credentials' contains:
          - file_path: path to the saved file on disk
          - filename: original filename
        """
        self._file_path = credentials.get("file_path")
        self._filename = credentials.get("filename")

        if not self._file_path or not os.path.exists(self._file_path):
            logger.error(f"File not found: {self._file_path}")
            return False

        ext = os.path.splitext(self._filename or "")[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            logger.error(f"Unsupported file extension: {ext}")
            return False

        return True

    async def test_connection(self) -> bool:
        return self._file_path is not None and os.path.exists(self._file_path)

    async def fetch_documents(self) -> List[NormalizedDocument]:
        """Extract text from the file and return as a NormalizedDocument."""
        if not self._file_path or not self._filename:
            raise ValueError("Connector not connected. Call connect() first.")

        ext = os.path.splitext(self._filename)[1].lower()
        content = ""

        try:
            if ext == ".txt":
                with open(self._file_path, "r", encoding="utf-8") as f:
                    content = f.read()

            elif ext == ".pdf":
                reader = PdfReader(self._file_path)
                parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        parts.append(text)
                content = "\n".join(parts)

            elif ext == ".docx":
                doc = DocxDocument(self._file_path)
                content = "\n".join(p.text for p in doc.paragraphs if p.text)

            else:
                raise ValueError(f"Unsupported file extension: {ext}")

        except Exception as e:
            logger.error(f"Failed to extract text from {self._filename}: {e}")
            raise ValueError(f"Text extraction failed: {e}")

        if not content.strip():
            raise ValueError("File content is empty or could not be extracted")

        doc_type = ext.lstrip(".")  # pdf, txt, docx

        return [
            NormalizedDocument(
                external_id=f"file-{uuid.uuid4()}",
                title=self._filename,
                content=content,
                doc_type=doc_type,
                metadata={
                    "file_path": self._file_path,
                    "file_size": os.path.getsize(self._file_path),
                }
            )
        ]
